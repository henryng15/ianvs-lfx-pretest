#!/usr/bin/env python3
"""Assemble submission/*.md plus every posted permalink into one .docx.

The email attachment takes precedence over GitHub for scoring, so this must
contain exactly what was posted. Bodies are read from the same files the posting
scripts used, and the link section is read from the recorded permalinks.
"""
import json, pathlib, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = pathlib.Path("submission/henryng15-lfx-ianvs-pretest.docx")
posted = json.loads(pathlib.Path("evidence/posted.json").read_text())
targets = json.loads(pathlib.Path("evidence/posted_targets.json").read_text())

doc = Document()
for s in ("Normal",):
    doc.styles[s].font.name = "Calibri"
    doc.styles[s].font.size = Pt(10.5)
doc.styles["Normal"].paragraph_format.space_after = Pt(6)


def mono(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x22, 0x33, 0x44)
    return p


INLINE = re.compile(r"`([^`]+)`|\*\*([^*]+)\*\*|\[([^\]]+)\]\(([^)]+)\)")


def prose(text):
    """Use a single ASCII hyphen in rendered prose.

    Code blocks bypass this helper so command flags and captured terminal output
    remain byte-for-byte faithful to the evidence.
    """
    return text.replace("—", "-").replace(" -- ", " - ")


def rich(par, text):
    """Render inline code, bold and links as runs."""
    text = prose(text)
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        if m.group(1) is not None:
            r = par.add_run(m.group(1)); r.font.name = "Consolas"; r.font.size = Pt(9)
        elif m.group(2) is not None:
            r = par.add_run(m.group(2)); r.bold = True
        else:
            r = par.add_run(m.group(3)); r.font.color.rgb = RGBColor(0x15, 0x5a, 0xb0)
            r.underline = True
            par.add_run(f" <{m.group(4)}>")
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def render(md):
    lines = md.splitlines()
    i, in_code, buf = 0, False, []
    while i < len(lines):
        ln = lines[i]
        if ln.startswith("```"):
            if in_code:
                mono("\n".join(buf)); buf = []; in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            buf.append(ln); i += 1; continue

        if ln.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1]):
            rows, j = [], i
            while j < len(lines) and lines[j].startswith("|"):
                rows.append([c.strip() for c in lines[j].strip("|").split("|")]); j += 1
            rows.pop(1)
            t = doc.add_table(rows=0, cols=max(len(r) for r in rows))
            t.style = "Light Grid Accent 1"
            for k, row in enumerate(rows):
                cells = t.add_row().cells
                for c, val in enumerate(row):
                    if c < len(cells):
                        cells[c].text = ""
                        p = cells[c].paragraphs[0]
                        p.paragraph_format.space_after = Pt(2)
                        rich(p, val)
                        for r in p.runs:
                            r.font.size = Pt(8.5)
                            if k == 0:
                                r.bold = True
            doc.add_paragraph()
            i = j; continue

        # a heading needs a space after the hashes -- "#642" is an issue reference
        hm = re.match(r"^(#{1,6})\s+(.*)$", ln)
        if hm:
            doc.add_heading(prose(hm.group(2).strip()), min(len(hm.group(1)), 4)); i += 1; continue
        if ln.strip() in ("---", "***"):
            doc.add_paragraph("_" * 78).alignment = WD_ALIGN_PARAGRAPH.CENTER; i += 1; continue
        if re.match(r"^\s*[-*]\s+", ln):
            p = doc.add_paragraph(style="List Bullet"); rich(p, re.sub(r"^\s*[-*]\s+", "", ln)); i += 1; continue
        if re.match(r"^\s*\d+\.\s+", ln):
            p = doc.add_paragraph(style="List Number"); rich(p, re.sub(r"^\s*\d+\.\s+", "", ln)); i += 1; continue
        if ln.startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            rich(p, ln.lstrip("> ")); 
            for r in p.runs: r.italic = True
            i += 1; continue
        if not ln.strip():
            i += 1; continue
        p = doc.add_paragraph(); rich(p, ln); i += 1


# ---- cover -----------------------------------------------------------------
doc.add_heading("LFX Mentorship 2026 Term 3 - Pre-test Submission", 0)
doc.add_heading("CNCF / KubeEdge - Comprehensive Example Restoration for KubeEdge Ianvs: Phase IV", 2)
for line in [
    "**Candidate:** henryng15 (Henry Nguyen)",
    "**Analysed commit:** kubeedge/ianvs@37a9c60",
    "**Submitted:** 2026-08-28",
]:
    p = doc.add_paragraph(); rich(p, line)

doc.add_heading("Submitted links", 1)
doc.add_heading("Pre-test Discussion", 3)
p = doc.add_paragraph(); rich(p, posted["discussion"]["url"])
doc.add_heading("Task comments", 3)
for label in ("Task 1", "Task 2", "Task 3", "Task 4", "Bonus"):
    p = doc.add_paragraph(style="List Bullet")
    rich(p, f"**{label}:** {posted['comments'][label]['url']}")

doc.add_heading("PR reviews (target-specific)", 3)
for k in ("pr-558", "pr-651", "pr-642", "pr-598", "pr-702"):
    p = doc.add_paragraph(style="List Bullet")
    rich(p, f"**Mandatory #{k[3:]}:** {targets[k]}")
for k in ("pr-598-followup", "pr-617", "pr-617-followup", "pr-617-correction",
          "pr-569", "pr-739", "pr-632", "pr-540"):
    p = doc.add_paragraph(style="List Bullet")
    label = k[3:].replace("-followup", " (follow-up)").replace("-correction", " (correction)")
    rich(p, f"**{'Mandatory' if k.startswith('pr-598') else 'Bonus'} #{label}:** {targets[k]}")

doc.add_heading("Issue comments (target-specific)", 3)
for k in ("issue-557", "issue-597", "issue-641", "issue-568"):
    p = doc.add_paragraph(style="List Bullet")
    rich(p, f"**#{k[6:]}:** {targets[k]}")

doc.add_heading("Reproduction repository", 3)
p = doc.add_paragraph(); rich(p, "https://github.com/henryng15/ianvs-lfx-pretest")

doc.add_page_break()

# ---- bodies ----------------------------------------------------------------
# each file already opens with its own H1, so no wrapper heading is added
for path in [
    "submission/00-discussion-body.md",
    "submission/01-task1.md",
    "submission/02-task2.md",
    "submission/03-task3.md",
    "submission/04-task4.md",
    "submission/05-bonus.md",
]:
    render(pathlib.Path(path).read_text())
    doc.add_page_break()

doc.add_heading("Appendix - target-specific comments as posted", 1)
order = [("PR", "558"), ("PR", "651"), ("PR", "642"), ("PR", "598"),
         ("PR", "598-followup"), ("PR", "702"), ("PR", "617"),
         ("PR", "617-followup"), ("PR", "617-correction"), ("PR", "569"), ("PR", "739"), ("PR", "632"),
         ("PR", "540")] + \
        [("Issue", str(i)) for i in (557, 597, 641, 568)]
for kind, num in order:
    key = f"{'pr' if kind == 'PR' else 'issue'}-{num}"
    label = num.replace("-followup", " (follow-up)").replace("-correction", " (correction)")
    doc.add_heading(f"{kind} #{label}", 2)
    p = doc.add_paragraph(); rich(p, targets[key])
    render(pathlib.Path(f"submission/reviews/{key}.md").read_text())

doc.save(OUT)
print(f"wrote {OUT}  ({OUT.stat().st_size/1024:.0f} KB)")
